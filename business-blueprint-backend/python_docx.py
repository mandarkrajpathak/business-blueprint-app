from docx import Document

def export_blueprint_to_docx(blueprint):
    doc = Document()
    doc.add_heading(f"Blueprint: {blueprint['project_name']}", level=1)

    for bp in blueprint["business_processes"]:
        doc.add_heading(bp["name"], level=2)
        doc.add_paragraph(f"Functional Requirement: {bp['functional_requirement']}")
        doc.add_paragraph(f"Integration Point: {bp['integration_point']}")
        doc.add_paragraph(f"RICEFW: {bp['ricefw']}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
