from docx import Document
import io

def generate_blueprint(parsed_requirements, project_name, module):
    for req in parsed_requirements:
        req["matched_by"] = "LLM"

    return {
        "project_name": project_name,
        "module": module,
        "sap_requirements": parsed_requirements,
        "non_sap_requirements": [],
        "bpmn_xml": None
    }

def export_blueprint_to_docx(blueprint: dict):
    doc = Document()
    doc.add_heading(f"Blueprint: {blueprint['project_name']}", level=1)
    doc.add_paragraph(f"Module: {blueprint['module']}")

    doc.add_heading("SAP Requirements", level=2)
    for req in blueprint.get("sap_requirements", []):
        doc.add_heading(req["business_process"], level=3)
        doc.add_paragraph(f"Functional Requirement: {req['functional_requirement']}")
        doc.add_paragraph(f"Suggested SAP Process: {req['sap_process_suggestion']} (Confidence: {req['sap_process_confidence']})")
        doc.add_paragraph(f"Explanation: {req['sap_process_explanation']}")
        doc.add_paragraph(f"Matched By: {req.get('matched_by', 'Unknown')}")
        if req["sap_process_roles"]:
            doc.add_paragraph("Roles: " + ", ".join(req["sap_process_roles"]))
        if req["sap_process_steps"]:
            doc.add_paragraph("Process Steps:")
            for step in req["sap_process_steps"]:
                doc.add_paragraph(f"• {step}", style='List Bullet')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
