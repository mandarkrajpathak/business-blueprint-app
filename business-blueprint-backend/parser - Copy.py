import io
from docx import Document

def parse_docx(file_bytes):
    document = Document(io.BytesIO(file_bytes))  # Wrap bytes for python-docx
    requirements = []

    current = {}
    print("Total paragraphs:", len(document.paragraphs))  # Debug

    for para in document.paragraphs:
        text = para.text.strip()
        print("Paragraph:", text)  # Debug

        if text.startswith("Business Process:"):
            current = {"business_process": text.split("Business Process:")[1].strip()}
        elif text.startswith("Functional Requirement:"):
            current["functional_requirement"] = text.split("Functional Requirement:")[1].strip()
        elif text.startswith("SAP Module:"):
            current["sap_module"] = text.split("SAP Module:")[1].strip()
        elif text.startswith("Integration Point:"):
            current["integration_point"] = text.split("Integration Point:")[1].strip()
        elif text.startswith("RICEFW:"):
            current["ricefw"] = text.split("RICEFW:")[1].strip()
            requirements.append(current)

    return requirements
